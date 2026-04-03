"""
Tạo file Word kịch bản thuyết trình chi tiết hoàn chỉnh
Gồm: FPS, khái niệm, cách đọc biểu đồ, quy trình chạy cụ thể, batch size
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ======================== PAGE SETTINGS ========================
section = doc.sections[0]
section.page_width  = Inches(8.27)   # A4
section.page_height = Inches(11.69)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ======================== STYLE HELPERS ========================
def add_title(doc, text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    run.font.size = Pt(18)

def add_h1(doc, text):
    p = doc.add_heading(text, level=1)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    run.font.size = Pt(14)

def add_h2(doc, text):
    p = doc.add_heading(text, level=2)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x0D, 0x6E, 0x6E)
    run.font.size = Pt(13)

def add_h3(doc, text):
    p = doc.add_heading(text, level=3)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
    run.font.size = Pt(12)

def add_script(doc, text):
    """Khung xanh – lời thoại trực tiếp"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("🎤 NÓI: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x0A, 0x60, 0x0A)
    run.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

def add_khainiem(doc, title, body):
    """Khung vàng – khái niệm kỹ thuật"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"📘 KHÁI NIỆM — {title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x7D, 0x5A, 0x00)
    run.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    r2.font.italic = True

def add_cachdoc(doc, title, body):
    """Khung tím – cách đọc biểu đồ"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(f"🖼️ CÁCH ĐỌC — {title}: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x5B, 0x2D, 0x8E)
    run.font.size = Pt(10)
    r2 = p.add_run(body)
    r2.font.size = Pt(10)

def add_hoidong(doc, question, answer):
    """Khung đỏ – câu hỏi hội đồng"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    r1 = p.add_run(f"❓ HỘI ĐỒNG HỎI: {question}\n")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)
    r1.font.size = Pt(10.5)
    r2 = p.add_run(f"✅ TRẢ LỜI: {answer}")
    r2.font.color.rgb = RGBColor(0x0A, 0x60, 0x0A)
    r2.font.size = Pt(10.5)

def add_bullet(doc, items):
    for item in items:
        p = doc.add_paragraph(style='List Bullet')
        if '**' in item:
            parts = item.split('**')
            for i, part in enumerate(parts):
                run = p.add_run(part)
                run.bold = (i % 2 == 1)
                run.font.size = Pt(11)
        else:
            run = p.add_run(item)
            run.font.size = Pt(11)

def add_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        tc = hdr[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A5376')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_text in enumerate(row_data):
            row[i].text = cell_text
            row[i].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()

def add_sep(doc):
    doc.add_paragraph("─" * 80)

# ================================================================
#  PHẦN 1: TIÊU ĐỀ
# ================================================================
add_title(doc, "🌾 KỊCH BẢN THUYẾT TRÌNH CHI TIẾT")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Đề tài: Phân loại Bệnh Lá Lúa bằng Deep Learning  |  Thời lượng: ~60 phút")
run.font.size = Pt(12)
run.italic = True

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Màu xanh lá = Lời thoại thực tế  |  Màu vàng = Khái niệm  |  Màu tím = Cách đọc hình  |  Màu đỏ = Câu hỏi hội đồng")
r2.font.size = Pt(9)
r2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
doc.add_paragraph()

# ================================================================
#  MỞ ĐẦU
# ================================================================
add_h1(doc, "🔰 MỞ ĐẦU (~3 phút)")
add_script(doc, 
    "Thưa thầy/cô, nhóm em xin trình bày đề tài Phân loại Bệnh Lá Lúa sử dụng Deep Learning tiên tiến. "
    "Việt Nam xuất khẩu gạo trong top đầu thế giới, nhưng thực tế đồng ruộng phần lớn nông dân vẫn nhận diện bệnh "
    "bằng kinh nghiệm mắt thường — rất chủ quan và chậm trễ. Nhóm em xây dựng hệ thống AI: "
    "chụp ảnh lá lúa → tự động phân loại 7 loại bệnh chỉ trong vài giây."
)
add_khainiem(doc, "Pipeline (Đường ống xử lý)",
    "Chuỗi 10 bước xử lý tuần tự, đầu ra bước trước thành đầu vào bước sau: "
    "Thu thập dữ liệu → Lọc & chuẩn hóa → EDA → Augmentation → DCGAN sinh ảnh → "
    "Chia tập → Huấn luyện → Đánh giá → Ensemble → Grad-CAM."
)
add_khainiem(doc, "GPU T4×2 / MirroredStrategy",
    "Kaggle cấp 2 card NVIDIA Tesla T4 (mỗi card 16 GB VRAM). MirroredStrategy sao chép "
    "mô hình lên cả 2 GPU, mỗi GPU xử lý một nửa batch, sau đó tổng hợp gradient (AllReduce) "
    "— huấn luyện nhanh ~gấp đôi."
)
add_khainiem(doc, "Batch Size",
    "Số ảnh được đưa vào mô hình cùng lúc trong 1 lần cập nhật trọng số. "
    "Nhóm dùng batch_size = 32: đủ lớn để gradient ổn định, đủ nhỏ để vừa VRAM 16 GB. "
    "Khi dùng 2 GPU (MirroredStrategy), mỗi GPU thực ra xử lý 16 ảnh/batch."
)

# ================================================================
#  STEP 1
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 1: HIỂU BÀI TOÁN & LỰA CHỌN MÔ HÌNH (~5 phút)")
add_script(doc,
    "Đây là bài toán Phân loại Đa lớp (Multi-class Classification) — đầu vào là ảnh RGB 224×224, "
    "đầu ra là 1 trong 7 nhãn bệnh. Nhóm chọn 4 mô hình CNN đại diện 4 trường phái "
    "từ siêu nhẹ đến cơ bắp nặng nhất."
)
add_h3(doc, "Bảng 4 mô hình CNN")
add_table(doc,
    ["Mô hình", "Năm", "Params", "Dung lượng", "FPS*", "Mục đích"],
    [
        ["MobileNetV3Large", "2019 (Google)", "~3.0 M", "~30 MB", "47 FPS", "Mobile / IoT"],
        ["EfficientNetV2S",  "2021 (Google)", "~21.5 M","~160 MB", "42 FPS", "Cân bằng tốc độ & độ chính xác"],
        ["ConvNeXtSmall",    "2022 (Meta)",   "~49.0 M","~380 MB", "35 FPS", "Sức mạnh CNN tối đa"],
        ["Proposed_SE",      "2026 (Nhóm)",   "~29.0 M","~255 MB", "39 FPS", "Đề xuất: Tiny+SE+FocalLoss"],
    ]
)
p = doc.add_paragraph()
r = p.add_run("(*) FPS đo trên GPU T4 inference đơn, batch_size = 1, ảnh 224×224. Ngưỡng Real-time ≥ 30 FPS.")
r.font.size = Pt(9)
r.font.italic = True

add_khainiem(doc, "SOTA (State Of The Art)",
    "Công nghệ tốt nhất thế giới tại thời điểm đó trong phân khúc CNN thuần túy. "
    "MobileNet = SOTA dòng siêu nhẹ (2019), EfficientNetV2 = SOTA tốc độ/chính xác (2021), "
    "ConvNeXt = SOTA CNN sức mạnh (2022). Tất cả đều là CNN, KHÔNG phải Transformer."
)
add_khainiem(doc, "Params (Tham số)",
    "Số lượng trọng số W và bias b mà mô hình cần học. "
    "Nhiều params → học được kiến thức phức tạp hơn nhưng cần nhiều VRAM và chạy chậm hơn."
)
add_khainiem(doc, "FPS (Frames Per Second — Khung hình trên giây)",
    "Số lượng ảnh mô hình có thể xử lý trong 1 giây khi chạy thực tế (inference). "
    "FPS cao = chạy nhanh hơn. Ngưỡng thực tế cần ≥ 30 FPS để camera cảm biến thời gian thực. "
    "MobileNet đạt 47 FPS → hoàn toàn chạy mượt trên smartphone yếu."
)

# ================================================================
#  STEP 2-3
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 2 & 3: DỮ LIỆU & EDA (~5 phút)")
add_script(doc,
    "Nhóm gộp 4 bộ dataset từ Kaggle thay vì 1 bộ đơn lẻ để tăng đa dạng điều kiện chụp. "
    "Sau đó lọc qua 3 bước: (1) Ảnh hỏng bằng PIL.verify(), (2) Ảnh quá nhỏ <32px, "
    "(3) Ảnh trùng nhau bằng MD5 hash."
)
add_khainiem(doc, "MD5 Hash",
    "Thuật toán tạo chuỗi 32 ký tự đại diện duy nhất (dấu vân tay) cho nội dung file. "
    "Hai ảnh pixel giống hệt → cùng hash → phát hiện trùng. Xác suất va chạm ≈ 1 / 2^128 ≈ 0."
)
add_h3(doc, "🖼️ Biểu đồ phân bố số lượng ảnh theo lớp")
add_cachdoc(doc, "Biểu đồ cột phân bố lớp",
    "Trục X = tên lớp bệnh, Trục Y = số ảnh. Cột cao nhất = Leaf Blast (~2041 ảnh), "
    "cột thấp nhất = Healthy (~523 ảnh). Chênh lệch gấp 4 lần → mất cân bằng nghiêm trọng "
    "→ đây là lý do dùng GAN và Focal Loss."
)
add_script(doc,
    "Thầy/cô nhìn vào biểu đồ cột sẽ thấy ngay 'cột lệch' rất rõ. "
    "Leaf Blast chiếm 27% tổng dữ liệu trong khi Healthy chỉ có 7%. "
    "Nếu dùng CrossEntropy thường, model sẽ bỏ quên lớp Healthy và chỉ cố phân loại đúng Leaf Blast."
)
add_h3(doc, "🖼️ Biểu đồ BoxPlot RGB")
add_cachdoc(doc, "BoxPlot RGB",
    "3 cột hộp (R, G, B) cho mỗi lớp bệnh. Hộp = 50% dữ liệu trung tâm. Đường giữa = Median. "
    "Râu = min/max. Điểm ngoài râu = outlier. "
    "Khi kênh G (Xanh lá) của Healthy và Hispa có hộp CHỒNG LÊN NHAU → "
    "2 lớp này không thể phân biệt bằng màu → CNN phải học đặc trưng hình dạng vết cắn."
)

# ================================================================
#  STEP 4
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 4: TIỀN XỬ LÝ & TĂNG CƯỜNG DỮ LIỆU (~7 phút)")
add_h3(doc, "4.1 Augmentation trên GPU")
add_script(doc,
    "Tất cả phép biến đổi chạy trực tiếp trên GPU bằng TensorFlow ops — "
    "không tốn thêm băng thông CPU-GPU. Mỗi ảnh được biến đổi ngẫu nhiên mỗi epoch."
)
add_table(doc,
    ["Phép biến đổi", "Tham số", "Tác dụng"],
    [
        ["Lật ngang/dọc", "50% xác suất", "Lá nhìn từ nhiều hướng"],
        ["Độ sáng",       "±20%",          "Mô phỏng điều kiện ánh sáng khác nhau"],
        ["Tương phản",    "0.8 – 1.2",     "Mô phỏng chất lượng camera khác nhau"],
        ["Bão hòa màu",   "0.8 – 1.2",     "Mô phỏng thời tiết / mùa vụ"],
        ["CutMix",        "λ ~ Beta(0.2)", "Cắt dán vùng bệnh giữa 2 ảnh khác lớp"],
        ["MixUp",         "λ ~ Beta(0.2)", "Pha trộn pixel + nhãn mềm (soft label)"],
    ]
)
add_khainiem(doc, "CutMix",
    "Cắt một vùng hình chữ nhật ngẫu nhiên từ ảnh B, dán đè lên ảnh A. "
    "Nhãn trộn theo tỷ lệ diện tích: nếu vùng cắt chiếm 30% → nhãn B chiếm 30%. "
    "Ép model nhận diện bệnh kể cả khi chỉ thấy một phần lá, tránh phụ thuộc phông nền."
)
add_khainiem(doc, "MixUp",
    "Ảnh_mới = λ × Ảnh_A + (1−λ) × Ảnh_B. Nhãn_mới = λ × Nhãn_A + (1−λ) × Nhãn_B. "
    "λ lấy từ phân phối Beta(0.2, 0.2) — dạng chữ U, thường λ gần 0 hoặc 1 → 1 ảnh chiếm ưu thế. "
    "Tạo ranh giới phân loại mượt hơn, giảm overconfidence."
)
add_khainiem(doc, "Batch Size trong Augmentation",
    "Augmentation CutMix/MixUp cần lấy 2 ảnh ngẫu nhiên TRONG cùng 1 batch. "
    "Batch_size = 32 → mỗi lần trộn sẽ chọn ngẫu nhiên cặp ảnh trong 32 ảnh đó. "
    "Batch càng lớn → nhiều cặp khả năng trộn → augmentation đa dạng hơn."
)
add_h3(doc, "4.2 CLAHE")
add_cachdoc(doc, "Trước/Sau CLAHE",
    "Bên trái: ảnh gốc bị lệch sáng, vết bệnh mờ. Bên phải: sau CLAHE kênh L trong không gian LAB, "
    "tương phản tăng, viền bệnh nổi bật. Màu sắc KHÔNG thay đổi vì chỉ chỉnh kênh Lightness."
)

# ================================================================
#  STEP 5
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 5: DCGAN SINH ẢNH CÂN BẰNG LỚP (~5 phút)")
add_script(doc,
    "Sau augmentation vẫn còn mất cân bằng, nhóm tự code DCGAN. "
    "Generator học vẽ ảnh lá bệnh giả từ nhiễu ngẫu nhiên. "
    "Discriminator làm 'giám khảo' phân biệt thật/giả. Hai bên cạnh tranh đến khi ảnh không thể phân biệt."
)
add_khainiem(doc, "Conditional DCGAN",
    "Thêm class_id (VD: 'Healthy') vào cả Generator và Discriminator để ép sinh đúng loại bệnh cần. "
    "Generator: noise(128D) → Dense(4×4×512) → 4 lớp ConvTranspose → ảnh 64×64×3. "
    "Discriminator: ảnh 64×64 → 4 lớp Conv → Dense(1) → xác suất thật/giả."
)
add_table(doc,
    ["Tham số", "Giá trị", "Ý nghĩa"],
    [
        ["noise_dim",  "128",    "Kích thước vector nhiễu đầu vào"],
        ["batch_size", "64",     "Số cặp ảnh thật/giả mỗi lần update — GAN cần batch lớn hơn CNN thường"],
        ["LR",         "2×10⁻⁴","Tốc độ học — GAN cần LR nhỏ để ổn định"],
        ["β₁ (Adam)",  "0.5",    "Momentum Adam — GAN dùng 0.5 thay vì 0.9 để giảm dao động"],
        ["steps",      "5000",   "Tổng số vòng cập nhật trọng số"],
    ]
)
add_khainiem(doc, "Batch Size trong GAN",
    "batch_size = 64 nghĩa là mỗi bước train: lấy 64 ảnh thật + sinh 64 ảnh giả → "
    "tổng 128 ảnh đưa vào Discriminator. Batch lớn hơn → gradient ổn định hơn → GAN ít dao động. "
    "Quy tắc vàng: TUYỆT ĐỐI không đưa ảnh GAN vào tập Validation và Test — chỉ Train!"
)

# ================================================================
#  STEP 6
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 6: HUẤN LUYỆN 3 PHASE (~8 phút)")
add_h3(doc, "6.1 Quy trình chạy chi tiết (từng bước cụ thể)")
add_script(doc,
    "Mỗi mô hình trong 4 mô hình đều trải qua đúng 3 giai đoạn (Phase) huấn luyện theo thứ tự bắt buộc. "
    "Không được đảo thứ tự."
)
add_table(doc,
    ["Phase", "Lớp trainable", "Learning Rate", "Epochs", "Batch Size", "Mục đích"],
    [
        ["1 — Warm-up",   "Chỉ Head",        "1e-3 (cố định)", "10 epochs", "32", "Dạy Head nhận diện 7 lớp mục tiêu"],
        ["2 — Fine-tune", "Head + 30% base", "Cosine 1e-4→0",  "20 epochs", "32", "Tinh chỉnh backbone cho lá lúa"],
        ["3 — Full fine", "100% tất cả lớp", "1e-5 (siêu nhỏ)","5 epochs",  "16", "Đánh bóng tinh tế, LR nhỏ bảo vệ backbone"],
    ]
)
add_khainiem(doc, "Batch Size thay đổi Phase 3",
    "Phase 3 giảm batch_size từ 32 xuống 16 vì mở toàn bộ lớp → bộ nhớ GPU tăng mạnh. "
    "Batch nhỏ hơn = gradient noisier nhưng giúp thoát local minima tốt hơn ở giai đoạn cuối."
)
add_khainiem(doc, "Freeze/Unfreeze (Đóng băng/Rã đông)",
    "Freeze (trainable=False): trọng số không thay đổi. Phase 1 đóng băng toàn bộ backbone → "
    "chỉ Head học → nhanh, không phá backbone. Phase 2 rã đông 30% backbone → "
    "tinh chỉnh đặc trưng cấp cao. Phase 3 rã đông 100% → đánh bóng tinh tế."
)
add_khainiem(doc, "Cosine Annealing LR",
    "LR(t) = LR_max × 0.5 × (1 + cos(π × t / T)). Đầu training: LR cao → hội tụ nhanh. "
    "Cuối training: LR gần 0 → tinh chỉnh rất mịn như 'hạ cánh nhẹ nhàng'. "
    "Tốt hơn constant LR vì giảm mượt, không bị plateau sớm."
)
add_khainiem(doc, "Callbacks (3 loại quan trọng)",
    "• EarlyStopping(patience=5): Val_loss không giảm 5 epoch liên tiếp → tự động dừng.\n"
    "• ReduceLROnPlateau(patience=2): Val_loss bình ổn 2 epoch → giảm LR×0.5.\n"
    "• ModelCheckpoint: Lưu trọng số TỐT NHẤT (val_loss thấp nhất) → không mất kết quả tốt."
)
add_h3(doc, "6.2 Mô hình Proposed: ConvNeXtTiny + SE-Block + Focal Loss")
add_script(doc,
    "Đây là đóng góp học thuật cốt lõi của nhóm. "
    "Không phải bưng nguyên xi mạng có sẵn mà tự ghép 3 thành phần lại với nhau."
)
add_khainiem(doc, "SE-Block (Squeeze-and-Excitation)",
    "Bước 1 — Squeeze: GlobalAveragePooling2D nén mỗi feature map (H×W) thành 1 số → vector C chiều.\n"
    "Bước 2 — Excitation: Dense(C/16, ReLU) → Dense(C, Sigmoid) → vector trọng số [0,1] mỗi kênh.\n"
    "Bước 3 — Scale: Nhân trọng số × feature map gốc. Kênh 'vết đốm nâu' → khuếch đại. Kênh 'trời xanh' → triệt tiêu.\n"
    "Chỉ tốn thêm ~500K params nhưng cải thiện đáng kể khả năng tập trung vào vùng bệnh."
)
add_khainiem(doc, "Focal Loss (γ=2.0, α=class_weight)",
    "Loss = −(1−p_t)^γ × log(p_t). Khi model đoán đúng dễ (p_t=0.9): (1−0.9)²=0.01 → loss ≈ 0 → bỏ qua. "
    "Khi model đoán sai (p_t=0.1): (1−0.1)²=0.81 → loss lớn → tập trung học. "
    "Giúp mô hình không bỏ quên lớp Healthy (chỉ 523 ảnh, 6.8% tổng)."
)
add_h3(doc, "🖼️ Biểu đồ Training Curves (4 mô hình)")
add_cachdoc(doc, "Training Curves",
    "Trục X = số epoch. Trục Y_trái = Accuracy, Trục Y_phải = Loss. "
    "Đường XANH = Training, Đường CAM = Validation. "
    "Đường thẳng đỏ dọc = Vạch Unfreeze (kết thúc Phase 1, bắt đầu Phase 2). "
    "Đường Validation nằm TRÊN Training là BÌNH THƯỜNG (do Dropout + CutMix làm Training khó hơn). "
    "Loss giảm đều = mô hình đang học tốt. Loss tăng trở lại = bắt đầu overfitting."
)
add_bullet(doc, [
    "**ConvNeXtSmall:** Học rất mượt, đường Loss hạ từ từ — kẻ lực điền đi bài bản.",
    "**EfficientNetV2S:** Sau vạch Unfreeze, Loss cắm đầu dốc xuống rất nhanh — hiệu quả nhất.",
    "**MobileNetV3Large:** Tại Epoch 12, Loss bị vọt lên khi Unfreeze — mạng quá nhỏ bị 'shock'. Sau tự phục hồi.",
    "**Proposed_SE:** Đường đẹp nhất: Loss Validation tụt cực sâu (<0.5), mượt, không gai. Focal Loss làm điều này.",
])

# ================================================================
#  STEP 7-8
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 7 & 8: ĐÁNH GIÁ — PHẦN CỨNG & ĐỘ CHÍNH XÁC (~10 phút)")
add_h3(doc, "🖼️ Hình 8: So sánh tài nguyên — Parameters (M), Model Size (MB), Inference Speed (FPS)")
add_cachdoc(doc, "Biểu đồ 3 panel Hardware",
    "Panel trái (màu hồng): Số tham số (Triệu). Panel giữa (màu xanh): Dung lượng trên đĩa (MB). "
    "Panel phải (màu xanh lá): Tốc độ suy luận FPS (Frames/giây). "
    "Đọc theo chiều: cột cao = nhiều hơn. Kết hợp cả 3 panel mới thấy rõ Trade-off. "
    "Model nhẹ (cột thấp panel trái) thường có FPS cao (cột cao panel phải) nhưng độ chính xác thấp hơn."
)
add_script(doc,
    "Thưa thầy/cô, trước khi xem ai thông minh hơn, hãy xem mỗi mô hình 'nặng bao nhiêu' và 'chạy nhanh bao nhiêu'. "
    "Đây là lý do nhóm phải thử nghiệm đủ 4 mô hình thay vì chỉ chọn 1."
)
add_h3(doc, "Bảng chi tiết 3 chỉ số phần cứng (đọc từ biểu đồ)")
add_table(doc,
    ["Mô hình", "Params (M)", "Dung lượng (MB)", "FPS (T4 GPU)", "Nhận xét"],
    [
        ["ConvNeXtSmall",    "~49 M", "~380 MB", "35 FPS", "Nặng nhất, chậm nhất — không phù hợp Mobile"],
        ["MobileNetV3Large", "~3 M",  "~30 MB",  "47 FPS", "Nhẹ nhất, nhanh nhất — lý tưởng Smartphone"],
        ["EfficientNetV2S",  "~21 M", "~160 MB", "42 FPS", "Cân bằng tốt — phù hợp App di động cấu hình trung"],
        ["Proposed_SE",      "~29 M", "~255 MB", "39 FPS", "≥30 FPS Real-time ✓ — dùng được camera thực tế"],
    ]
)
add_khainiem(doc, "FPS Real-time threshold",
    "Tiêu chuẩn quốc tế: ≥ 30 FPS = Real-time (thời gian thực) cho camera xử lý ảnh. "
    "Proposed đạt 39 FPS → VƯỢT ngưỡng Real-time → có thể nhúng vào camera giám sát ruộng. "
    "ConvNeXtSmall chỉ 35 FPS → vẫn đạt nhưng sát ngưỡng hơn."
)
add_hoidong(doc,
    "Tại sao Proposed 29 triệu params mà file lại nặng tới 255 MB? Còn MobileNet 3 triệu mà chỉ 30 MB. Tỷ lệ lệch quá!",
    "Dạ thưa, file H5 khi lưu ModelCheckpoint tự động đính kèm Trạng thái Tối ưu hóa (Optimizer State) của hàm Adam. "
    "Adam cần lưu thêm 2 biến: Momentum và Variance cho TỪNG tham số. "
    "28.6M × 3 = 85.8M biến × 4 bytes (float32) ≈ 343 MB. "
    "MobileNet chỉ 3M params → 3M×3×4 ≈ 36 MB. Khi xuất thực tế, bỏ optimizer đi "
    "thì Proposed chỉ còn 29M×4 bytes ≈ 116 MB thôi ạ."
)

doc.add_paragraph()
add_h3(doc, "🖼️ Biểu đồ So sánh 6 chỉ số độ chính xác Test Set")
add_cachdoc(doc, "Bar Chart 6 chỉ số",
    "Trục X = tên mô hình. Trục Y = giá trị (0 đến 1). 6 nhóm cột màu = 6 chỉ số khác nhau. "
    "Cột cao = tốt hơn (trừ Loss). Đọc theo nhóm màu: xem cùng 1 màu để so sánh cùng chỉ số giữa 4 model. "
    "Đường giá trị 0.88, 0.87... — chênh 1-2% trong thực tế không có ý nghĩa thống kê."
)
add_table(doc,
    ["Chỉ số", "EfficientNetV2S", "ConvNeXtSmall", "Proposed_SE", "MobileNetV3L", "Model tốt nhất"],
    [
        ["Accuracy",    "0.888 ★", "0.886", "0.876", "0.857", "EfficientNet"],
        ["Precision",   "0.871",   "0.872 ★","0.857","0.842", "ConvNeXt"],
        ["Recall",      "0.889 ★", "0.873", "0.872", "0.845", "EfficientNet"],
        ["F1-Macro",    "0.878 ★", "0.873", "0.861", "0.860", "EfficientNet"],
        ["F1-Weighted", "0.891 ★", "0.887", "0.881", "0.879", "EfficientNet"],
        ["Kappa",       "0.866 ★", "0.862", "0.852", "0.829", "Tất cả > 0.81 ✓"],
    ]
)
add_khainiem(doc, "Accuracy",  "Tỷ lệ đúng / tổng số mẫu. Đơn giản nhưng không đủ khi dữ liệu mất cân bằng.")
add_khainiem(doc, "Recall (Sensitivity)", 
    "Trong tất cả mẫu THỰC SỰ bị bệnh X, model phát hiện được bao nhiêu %. "
    "Trong nông nghiệp: Recall QUAN TRỌNG HƠN Precision vì bỏ sót bệnh → lây lan toàn đồng.")
add_khainiem(doc, "F1-Macro",
    "Trung bình F1 từng lớp, MỌI lớp tính bằng nhau. "
    "Chỉ số công bằng nhất khi dữ liệu mất cân bằng — EfficientNet đạt cao nhất 0.878.")
add_khainiem(doc, "Cohen's Kappa",
    "Đo đồng thuận giữa model và ground truth, ĐÃ LOẠI BỎ yếu tố may mắn. "
    "κ > 0.81 = 'Almost Perfect Agreement'. CẢ 4 model đều đạt → không model nào ăn may.")

# ================================================================
#  STEP 9: ENSEMBLE
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 9: QUY TRÌNH INFERENCE THỰC TẾ — ENSEMBLE SOFT-VOTING (~5 phút)")
add_h3(doc, "Quy trình chạy chi tiết từng bước khi người dùng tải ảnh lên:")
add_table(doc,
    ["Bước", "Mô tả chi tiết", "Đầu ra"],
    [
        ["1. Tiền xử lý", "Resize ảnh về 224×224, normalize pixel [0,1], áp dụng CLAHE", "Tensor 1×224×224×3"],
        ["2. TTA (×4)",   "Tạo 4 biến thể: gốc + lật ngang + lật dọc + lật cả 2", "4 tensor đầu vào"],
        ["3. Model 1",    "ConvNeXtSmall dự đoán 4 biến thể → trung bình → vector 7 xác suất", "[0.05, 0.80, ...]"],
        ["4. Model 2",    "MobileNetV3Large dự đoán → vector 7 xác suất", "[0.03, 0.75, ...]"],
        ["5. Model 3",    "EfficientNetV2S dự đoán → vector 7 xác suất", "[0.02, 0.85, ...]"],
        ["6. Model 4",    "Proposed_SE dự đoán → vector 7 xác suất", "[0.04, 0.78, ...]"],
        ["7. Soft-Voting", "Cộng 4 vector → chia 4 → lấy argmax", "Bệnh: Brown_Spot 79.5%"],
        ["8. Grad-CAM",   "Gọi Grad-CAM cho từng model với nhãn vừa chọn → 4 ảnh nhiệt", "4 heatmap overlay"],
        ["9. Xuất kết quả", "Hiển thị: bệnh + % + 4 ảnh Grad-CAM giải thích", "Kết quả cuối"],
    ]
)
add_khainiem(doc, "TTA (Test-Time Augmentation)",
    "Mỗi ảnh được dự đoán trên 4 biến thể (gốc + 3 flip) → trung bình xác suất 4 lần → chọn max. "
    "Tăng accuracy ~0.5-1% mà không cần train lại. Đánh đổi: inference chậm hơn ~4 lần."
)
add_khainiem(doc, "Soft-Voting Ensemble",
    "KHÔNG phải Boosting (XGBoost — học tuần tự sửa lỗi). "
    "Ensemble CNN: 4 mô hình chạy SONG SONG ĐỘC LẬP, mỗi model xuất vector xác suất, "
    "cộng lại chia trung bình, chọn lớp cao nhất. Giống 'hội chẩn bác sĩ' thay vì 'một người quyết.'"
)
add_khainiem(doc, "Grad-CAM KHÔNG chạy trên Ensemble",
    "Ensemble là phép cộng trung bình — không có lớp Conv chung → KHÔNG THỂ tính Grad-CAM. "
    "Grad-CAM chạy ĐỘC LẬP trên từng mô hình RIÊNG BIỆT sau khi Ensemble đã chốt nhãn. "
    "Kết quả: 4 ảnh nhiệt riêng biệt cho mỗi mô hình — minh bạch hơn 1 ảnh tổng hợp."
)

# ================================================================
#  STEP 10: GRAD-CAM
# ================================================================
add_sep(doc)
add_h1(doc, "STEP 10: GRAD-CAM — CHỨNG MINH AI KHÔNG ĐOÁN LỤI (~5 phút)")
add_h3(doc, "🖼️ Hình Grad-CAM 4 mô hình (5 cột × N hàng)")
add_cachdoc(doc, "Hình Grad-CAM 4 Model",
    "Cột 1 = Ảnh gốc. Cột 2-5 = Heatmap overlay của 4 model theo thứ tự. "
    "Màu ĐỎ/VÀNG = model tập trung vào vùng đó để ra quyết định. "
    "Màu XANH LẠNH = model bỏ qua vùng đó (ít thông tin). "
    "Đọc: nếu vùng đỏ trùng với vết bệnh thật trên lá → model đang học ĐÚNG đặc trưng bệnh."
)
add_script(doc,
    "Thầy/cô nhìn vào hình này, theo từng hàng (cùng 1 ảnh lá): "
    "(1) ConvNeXtSmall: vùng đỏ lan rộng — mạng lớn nhìn 'tổng quan' toàn bộ lá trước khi quyết. "
    "(2) MobileNetV3: quầng đỏ hẹp hơn, đôi khi lệch viền — bị giới hạn bộ nhớ không gian do ít params. "
    "(3) EfficientNetV2S: vùng đỏ chính xác, cắm chặt vào vết bệnh. "
    "(4) Proposed_SE: nhờ SE-Block lọc kênh, vùng đỏ SẮC NÉT và TẬP TRUNG NHẤT — "
    "bỏ qua hoàn toàn phông nền, bầu trời, đất đá. Minh chứng SE-Block hoạt động hiệu quả."
)
add_khainiem(doc, "Grad-CAM (Gradient-weighted Class Activation Mapping)",
    "Bước 1: Forward pass ảnh → lấy logit lớp dự đoán.\n"
    "Bước 2: Tính gradient của logit đó theo feature map của lớp Conv2D cuối cùng.\n"
    "Bước 3: Global Average Pooling gradient → trọng số (importance) mỗi kênh.\n"
    "Bước 4: Nhân trọng số × feature map → cộng tất cả kênh → ReLU → heatmap thô.\n"
    "Bước 5: Resize heatmap về 224×224 (bicubic) → overlay lên ảnh gốc với alpha=0.5."
)

# ================================================================
#  TỔNG KẾT
# ================================================================
add_sep(doc)
add_h1(doc, "TỔNG KẾT (~2 phút)")
add_script(doc,
    "Tóm lại, nhóm xây dựng pipeline 10 bước hoàn chỉnh. EfficientNetV2S dẫn đầu 88.8% Accuracy. "
    "Mô hình Proposed của nhóm đạt 87.6% — chỉ kém 1.2% dù backbone chỉ là ConvNeXtTiny, "
    "đồng thời đạt 39 FPS — vượt ngưỡng Real-time 30 FPS cho ứng dụng camera thực tế. "
    "Toàn bộ pipeline sử dụng CNN thuần túy — không có Transformer. "
    "Nhóm em xin cảm ơn thầy/cô và kính mời hội đồng đặt câu hỏi!"
)

# ================================================================
#  PHỤ LỤC
# ================================================================
add_sep(doc)
add_h1(doc, "PHỤ LỤC: CÂU HỎI HỘI ĐỒNG & TRẢ LỜI MẪU")
add_hoidong(doc,
    "Tại sao Train Accuracy lại THẤP HƠN Validation Accuracy? Nghe kỳ lạ quá!",
    "Bình thường ạ. Khi train: Dropout tắt ngẫu nhiên 30% neuron, MixUp/CutMix trộn ảnh lộn xộn, "
    "Label Smoothing làm nhãn mềm → training KHÓ HƠN. Khi validate/test: tắt hết augmentation, "
    "Dropout = 0 → dùng full năng lực → điểm cao hơn. Đây là dấu hiệu regularization hoạt động tốt."
)
doc.add_paragraph()
add_hoidong(doc,
    "ViT và Swin Transformer bây giờ mới là SOTA, sao không dùng?",
    "ViT là SOTA trên siêu máy chủ ImageNet. ConvNeXt (2022) sinh ra để chứng minh CNN thuần túy "
    "vẫn cạnh tranh ngang ViT. Hơn nữa, ViT cần rất nhiều data (~300M ảnh) mà nhóm chỉ có ~10K. "
    "Và quan trọng nhất: ViT inference chậm — không đạt 30+ FPS trên thiết bị nhỏ. CNN là lựa chọn thực tế."
)
doc.add_paragraph()
add_hoidong(doc,
    "Ensemble rồi, Grad-CAM chạy trên ensemble hay từng model?",
    "Grad-CAM chạy ĐỘC LẬP trên từng model sau khi Ensemble đã chốt nhãn. "
    "Lý do: Ensemble là phép cộng trung bình — không có lớp Conv chung để tính gradient. "
    "Kết quả là 4 heatmap riêng biệt — minh bạch hơn, cho thấy mỗi model 'nhìn' chỗ nào khác nhau."
)
doc.add_paragraph()
add_hoidong(doc,
    "Proposed 29M params gọi là Tiny nhưng nặng hơn EfficientNetV2S 21M. Logic nào vậy?",
    "ConvNeXt-Tiny là bản NHỎ NHẤT trong gia đình ConvNeXt (Tiny < Small < Base < Large < XLarge). "
    "ConvNeXt sinh ra để cạnh tranh Transformer nên ngay bản Tiny cũng đã nặng 28M. "
    "SE-Block nhóm thêm vào chỉ tốn thêm ~500K params. "
    "ConvNeXt-Tiny NẶNG HƠN EfficientNetV2S vì 2 họ kiến trúc hoàn toàn khác nhau — "
    "ConvNeXt dùng kernel 7×7 và nhiều kênh hơn."
)
doc.add_paragraph()
add_hoidong(doc,
    "Tại sao Proposed thấp hơn EfficientNet? SE-Block không giúp ích à?",
    "3 lý do: (1) Backbone Tiny có 28M params vs EfficientNet 21M nhưng kiến trúc EfficientNet "
    "được Google NAS tối ưu hoàn hảo hơn. "
    "(2) Focal Loss hy sinh một chút Precision lớp đa số để tăng Recall lớp thiểu số — đánh đổi có chủ đích. "
    "(3) Khoảng cách 1.2% ≈ 14 ảnh trên 1156 ảnh — không có ý nghĩa thống kê. "
    "Nếu ghép SE-Block vào ConvNeXtSmall thay vì Tiny, khả năng rất cao sẽ vượt EfficientNet."
)
doc.add_paragraph()
add_hoidong(doc,
    "Ensemble là Boosting hay Voting? Có phải Bagging không?",
    "Soft-Voting Ensemble. Không phải Boosting (học tuần tự, con sau sửa lỗi con trước — như XGBoost). "
    "Không phải Bagging (train nhiều model trên subset data ngẫu nhiên — như Random Forest). "
    "Ensemble nhóm: 4 model train ĐỘC LẬP trên TOÀN BỘ data, sau đó cộng trung bình xác suất."
)

# ================================================================
#  SAVE
# ================================================================
out_path = r"D:\DataDownloaded\Study\MonHoc\Do_An\deeplearning\CK2\Kich_Ban_Bao_Cao_HoanChinh_V4.docx"
doc.save(out_path)
print(f"✅ Đã xuất file Word: {out_path}")
